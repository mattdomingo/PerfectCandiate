import os, io, re, datetime
from typing import Dict, Any
from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML
from docx import Document
from docx.shared import Pt

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
_env = Environment(
    loader=FileSystemLoader(TEMPLATES_DIR),
    autoescape=select_autoescape(["html", "xml"])
)

def _slug(s: str, fallback: str = "resume"):
    if not s: return fallback
    s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")
    return s or fallback

def render_pdf(json_resume: Dict[str, Any]) -> bytes:
    tpl = _env.get_template("resume.html.j2")
    html = tpl.render(**json_resume)
    pdf_bytes = HTML(string=html).write_pdf()
    return pdf_bytes

def render_docx(json_resume: Dict[str, Any]) -> bytes:
    # build a simple .docx without an external template
    doc = Document()
    styles = doc.styles
    try:
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)
    except Exception:
        pass

    basics = json_resume.get("basics", {}) or {}
    work   = json_resume.get("work", []) or []
    skills = json_resume.get("skills", []) or []
    edu    = json_resume.get("education", []) or []

    # Name
    name = basics.get("name") or "Your Name"
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)

    # meta line
    parts = []
    for k in ("email","phone","location"):
        v = basics.get(k)
        if v: parts.append(v)
    if parts:
        doc.add_paragraph(" • ".join(parts))

    # Experience
    if work:
        doc.add_paragraph().add_run("Experience").bold = True
        for w in work:
            header = []
            if w.get("position"): header.append(str(w["position"]))
            if w.get("company"): header.append(str(w["company"]))
            dates = []
            if w.get("start"): dates.append(str(w["start"]))
            if w.get("end"):   dates.append(str(w["end"]))
            elif w.get("start"): dates.append("Present")
            if dates:
                header.append("(" + " – ".join(dates) + ")")
            doc.add_paragraph(" — ".join(header))
            for h in (w.get("highlights") or []):
                doc.add_paragraph(h, style=None).style = doc.styles["List Bullet"]

    # Skills
    if skills:
        doc.add_paragraph().add_run("Skills").bold = True
        line = []
        for s in skills:
            chunk = s.get("name","")
            kws = s.get("keywords") or []
            if kws:
                chunk += ": " + ", ".join(map(str, kws))
            line.append(chunk)
        doc.add_paragraph("; ".join(line))

    # Education
    if edu:
        doc.add_paragraph().add_run("Education").bold = True
        for e in edu:
            line = []
            if e.get("degree"): line.append(str(e["degree"]))
            if e.get("institution"): line.append(str(e["institution"]))
            if e.get("start") or e.get("end"):
                line.append(f"({e.get('start','')} – {e.get('end','')})")
            doc.add_paragraph(", ".join(line))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def filename_for(json_resume: Dict[str, Any], ext: str) -> str:
    name = _slug((json_resume.get("basics") or {}).get("name") or "resume")
    stamp = datetime.datetime.utcnow().strftime("%Y%m%d")
    return f"{name}_{stamp}.{ext}"


