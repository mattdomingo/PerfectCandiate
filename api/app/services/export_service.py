"""
Export service for generating PDF and DOCX resumes
"""
import os
import tempfile
from typing import Dict, Any, Optional, Tuple, List
import logging
from datetime import datetime
import json

# Document generation imports
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template, Environment, FileSystemLoader

from app.services.storage import StorageService

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting resumes in various formats"""
    
    def __init__(self):
        self.storage_service = StorageService()
        
        # Setup Jinja2 environment for templates
        template_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
        os.makedirs(template_dir, exist_ok=True)
        
        self.jinja_env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=True
        )
        
        # Ensure template exists
        self._ensure_default_template()
    
    async def export_resume_pdf(self, resume_data: Dict[str, Any], suggestions: Optional[List[Dict[str, Any]]] = None) -> Tuple[str, bytes]:
        """
        Export resume as PDF
        
        Args:
            resume_data: Structured resume data
            suggestions: Optional list of applied suggestions
            
        Returns:
            Tuple of (filename, pdf_bytes)
        """
        try:
            # Generate HTML from template
            html_content = await self._generate_html(resume_data, suggestions)
            
            # Generate PDF using WeasyPrint
            pdf_bytes = await self._html_to_pdf(html_content)
            
            # Generate filename
            name = resume_data.get('contact_info', {}).get('name', 'resume')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{name.replace(' ', '_')}_{timestamp}.pdf"
            
            return filename, pdf_bytes
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise
    
    async def export_resume_docx(self, resume_data: Dict[str, Any], suggestions: Optional[List[Dict[str, Any]]] = None) -> Tuple[str, bytes]:
        """
        Export resume as DOCX
        
        Args:
            resume_data: Structured resume data
            suggestions: Optional list of applied suggestions
            
        Returns:
            Tuple of (filename, docx_bytes)
        """
        try:
            # Generate DOCX document
            docx_bytes = await self._generate_docx(resume_data, suggestions)
            
            # Generate filename
            name = resume_data.get('contact_info', {}).get('name', 'resume')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{name.replace(' ', '_')}_{timestamp}.docx"
            
            return filename, docx_bytes
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    async def save_exported_resume(self, resume_id: int, format_type: str, filename: str, content: bytes) -> str:
        """
        Save exported resume to storage
        
        Args:
            resume_id: Resume ID
            format_type: 'pdf' or 'docx'
            filename: Export filename
            content: File content as bytes
            
        Returns:
            Storage object name
        """
        try:
            folder = f"exports/{resume_id}"
            content_type = "application/pdf" if format_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            object_name = await self.storage_service.upload_file(
                file_data=content,
                filename=filename,
                content_type=content_type,
                folder=folder
            )
            
            return object_name
            
        except Exception as e:
            logger.error(f"Error saving exported resume: {e}")
            raise
    
    async def _generate_html(self, resume_data: Dict[str, Any], suggestions: Optional[List[Dict[str, Any]]] = None) -> str:
        """Generate HTML content from resume data"""
        try:
            template = self.jinja_env.get_template('resume.html')
            
            # Prepare template data
            template_data = {
                'resume': resume_data,
                'contact_info': resume_data.get('contact_info', {}),
                'work_experience': resume_data.get('work_experience', []),
                'education': resume_data.get('education', []),
                'skills': resume_data.get('skills', []),
                'certifications': resume_data.get('certifications', []),
                'summary': resume_data.get('summary', ''),
                'projects': resume_data.get('projects', []),
                'applied_suggestions': suggestions or [],
                'export_date': datetime.now().strftime("%B %d, %Y")
            }
            
            html_content = template.render(**template_data)
            return html_content
            
        except Exception as e:
            logger.error(f"Error generating HTML: {e}")
            raise
    
    async def _html_to_pdf(self, html_content: str) -> bytes:
        """Convert HTML to PDF using WeasyPrint"""
        try:
            # CSS for better PDF styling
            css_content = """
            @page {
                size: A4;
                margin: 0.75in;
            }
            body {
                font-family: 'Arial', sans-serif;
                line-height: 1.4;
                color: #333;
            }
            .header {
                text-align: center;
                border-bottom: 2px solid #333;
                padding-bottom: 10px;
                margin-bottom: 20px;
            }
            .section {
                margin-bottom: 20px;
            }
            .section-title {
                font-size: 16px;
                font-weight: bold;
                color: #2c3e50;
                border-bottom: 1px solid #bdc3c7;
                padding-bottom: 2px;
                margin-bottom: 10px;
            }
            .experience-item, .education-item {
                margin-bottom: 15px;
            }
            .job-title, .degree {
                font-weight: bold;
                color: #2c3e50;
            }
            .company, .school {
                font-style: italic;
                color: #7f8c8d;
            }
            .dates {
                float: right;
                color: #7f8c8d;
            }
            .skills-list {
                display: flex;
                flex-wrap: wrap;
            }
            .skill-item {
                background-color: #ecf0f1;
                padding: 4px 8px;
                margin: 2px;
                border-radius: 4px;
                font-size: 12px;
            }
            """
            
            css = CSS(string=css_content)
            html = HTML(string=html_content)
            
            pdf_bytes = html.write_pdf(stylesheets=[css])
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"Error converting HTML to PDF: {e}")
            raise
    
    async def _generate_docx(self, resume_data: Dict[str, Any], suggestions: Optional[List[Dict[str, Any]]] = None) -> bytes:
        """Generate DOCX document from resume data"""
        try:
            doc = Document()
            
            # Set up styles
            self._setup_docx_styles(doc)
            
            # Add header with contact info
            self._add_contact_header(doc, resume_data.get('contact_info', {}))
            
            # Add summary if available
            summary = resume_data.get('summary', '')
            if summary:
                self._add_section(doc, "Professional Summary", summary)
            
            # Add work experience
            work_exp = resume_data.get('work_experience', [])
            if work_exp:
                self._add_experience_section(doc, work_exp)
            
            # Add education
            education = resume_data.get('education', [])
            if education:
                self._add_education_section(doc, education)
            
            # Add skills
            skills = resume_data.get('skills', [])
            if skills:
                self._add_skills_section(doc, skills)
            
            # Add certifications
            certifications = resume_data.get('certifications', [])
            if certifications:
                self._add_certifications_section(doc, certifications)
            
            # Add projects if available
            projects = resume_data.get('projects', [])
            if projects:
                self._add_projects_section(doc, projects)
            
            # Save to bytes
            with tempfile.NamedTemporaryFile() as tmp:
                doc.save(tmp.name)
                tmp.seek(0)
                return tmp.read()
                
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    def _setup_docx_styles(self, doc):
        """Setup custom styles for DOCX document"""
        styles = doc.styles
        
        # Title style
        if 'Resume Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Resume Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section heading style
        if 'Section Heading' not in [style.name for style in styles]:
            heading_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = None  # Default color
    
    def _add_contact_header(self, doc, contact_info: Dict[str, Any]):
        """Add contact information header"""
        # Name
        name = contact_info.get('name', 'Your Name')
        name_para = doc.add_paragraph(name)
        name_para.style = 'Resume Title'
        
        # Contact details
        contact_details = []
        if contact_info.get('email'):
            contact_details.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_details.append(contact_info['phone'])
        if contact_info.get('location'):
            contact_details.append(contact_info['location'])
        
        if contact_details:
            contact_para = doc.add_paragraph(' | '.join(contact_details))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_section(self, doc, title: str, content: str):
        """Add a general section with title and content"""
        # Section title
        title_para = doc.add_paragraph(title)
        title_para.style = 'Section Heading'
        
        # Section content
        doc.add_paragraph(content)
        doc.add_paragraph()  # Spacing
    
    def _add_experience_section(self, doc, work_experience: List[Dict[str, Any]]):
        """Add work experience section"""
        title_para = doc.add_paragraph("Professional Experience")
        title_para.style = 'Section Heading'
        
        for exp in work_experience:
            # Job title and company
            job_line = f"{exp.get('title', '')} - {exp.get('company', '')}"
            if exp.get('location'):
                job_line += f" | {exp['location']}"
            
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(job_line)
            job_run.bold = True
            
            # Dates
            dates = ""
            if exp.get('start_date') and exp.get('end_date'):
                dates = f"{exp['start_date']} - {exp['end_date']}"
            elif exp.get('start_date'):
                dates = f"{exp['start_date']} - Present"
            
            if dates:
                date_para = doc.add_paragraph(dates)
                date_para.style.font.italic = True
            
            # Responsibilities
            responsibilities = exp.get('responsibilities', [])
            if responsibilities:
                for resp in responsibilities:
                    bullet_para = doc.add_paragraph(resp, style='List Bullet')
            
            doc.add_paragraph()  # Spacing between jobs
    
    def _add_education_section(self, doc, education: List[Dict[str, Any]]):
        """Add education section"""
        title_para = doc.add_paragraph("Education")
        title_para.style = 'Section Heading'
        
        for edu in education:
            # Degree and field
            degree_line = f"{edu.get('degree', '')} in {edu.get('field', '')}"
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(degree_line)
            degree_run.bold = True
            
            # School and location
            school_line = edu.get('school', '')
            if edu.get('location'):
                school_line += f" | {edu['location']}"
            
            if school_line:
                school_para = doc.add_paragraph(school_line)
                school_para.style.font.italic = True
            
            # Graduation date and GPA
            details = []
            if edu.get('graduation_date'):
                details.append(f"Graduated: {edu['graduation_date']}")
            if edu.get('gpa'):
                details.append(f"GPA: {edu['gpa']}")
            
            if details:
                doc.add_paragraph(' | '.join(details))
            
            doc.add_paragraph()  # Spacing
    
    def _add_skills_section(self, doc, skills: List[str]):
        """Add skills section"""
        title_para = doc.add_paragraph("Skills")
        title_para.style = 'Section Heading'
        
        # Group skills into lines for better formatting
        skills_text = ', '.join(skills)
        doc.add_paragraph(skills_text)
        doc.add_paragraph()  # Spacing
    
    def _add_certifications_section(self, doc, certifications: List[Dict[str, Any]]):
        """Add certifications section"""
        title_para = doc.add_paragraph("Certifications")
        title_para.style = 'Section Heading'
        
        for cert in certifications:
            cert_line = cert.get('name', '')
            if cert.get('issuer'):
                cert_line += f" - {cert['issuer']}"
            if cert.get('date'):
                cert_line += f" ({cert['date']})"
            
            doc.add_paragraph(cert_line, style='List Bullet')
        
        doc.add_paragraph()  # Spacing
    
    def _add_projects_section(self, doc, projects: List[Dict[str, Any]]):
        """Add projects section"""
        title_para = doc.add_paragraph("Projects")
        title_para.style = 'Section Heading'
        
        for project in projects:
            # Project name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(project.get('name', ''))
            name_run.bold = True
            
            # Project description
            description = project.get('description', '')
            if description:
                doc.add_paragraph(description)
            
            # Technologies
            technologies = project.get('technologies', [])
            if technologies:
                tech_text = f"Technologies: {', '.join(technologies)}"
                tech_para = doc.add_paragraph(tech_text)
                tech_para.style.font.italic = True
            
            doc.add_paragraph()  # Spacing
    
    def _ensure_default_template(self):
        """Ensure default HTML template exists"""
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'resume.html')
        
        if not os.path.exists(template_path):
            os.makedirs(os.path.dirname(template_path), exist_ok=True)
            
            default_template = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume - {{ contact_info.name or 'Your Name' }}</title>
</head>
<body>
    <div class="header">
        <h1>{{ contact_info.name or 'Your Name' }}</h1>
        <div class="contact-info">
            {% if contact_info.email %}<span>{{ contact_info.email }}</span>{% endif %}
            {% if contact_info.phone %}<span>{{ contact_info.phone }}</span>{% endif %}
            {% if contact_info.location %}<span>{{ contact_info.location }}</span>{% endif %}
        </div>
    </div>

    {% if summary %}
    <div class="section">
        <h2 class="section-title">Professional Summary</h2>
        <p>{{ summary }}</p>
    </div>
    {% endif %}

    {% if work_experience %}
    <div class="section">
        <h2 class="section-title">Professional Experience</h2>
        {% for exp in work_experience %}
        <div class="experience-item">
            <div class="job-header">
                <span class="job-title">{{ exp.title or '' }}</span>
                {% if exp.company %}<span class="company"> - {{ exp.company }}</span>{% endif %}
                {% if exp.start_date or exp.end_date %}
                <span class="dates">{{ exp.start_date or '' }}{% if exp.end_date %} - {{ exp.end_date }}{% else %} - Present{% endif %}</span>
                {% endif %}
            </div>
            {% if exp.responsibilities %}
            <ul>
                {% for resp in exp.responsibilities %}
                <li>{{ resp }}</li>
                {% endfor %}
            </ul>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if education %}
    <div class="section">
        <h2 class="section-title">Education</h2>
        {% for edu in education %}
        <div class="education-item">
            <span class="degree">{{ edu.degree or '' }}{% if edu.field %} in {{ edu.field }}{% endif %}</span>
            {% if edu.school %}<span class="school"> - {{ edu.school }}</span>{% endif %}
            {% if edu.graduation_date %}<span class="dates">{{ edu.graduation_date }}</span>{% endif %}
            {% if edu.gpa %}<div class="gpa">GPA: {{ edu.gpa }}</div>{% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if skills %}
    <div class="section">
        <h2 class="section-title">Skills</h2>
        <div class="skills-list">
            {% for skill in skills %}
            <span class="skill-item">{{ skill }}</span>
            {% endfor %}
        </div>
    </div>
    {% endif %}

    {% if certifications %}
    <div class="section">
        <h2 class="section-title">Certifications</h2>
        {% for cert in certifications %}
        <div class="cert-item">
            {{ cert.name or '' }}{% if cert.issuer %} - {{ cert.issuer }}{% endif %}{% if cert.date %} ({{ cert.date }}){% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if projects %}
    <div class="section">
        <h2 class="section-title">Projects</h2>
        {% for project in projects %}
        <div class="project-item">
            <h3>{{ project.name or '' }}</h3>
            {% if project.description %}<p>{{ project.description }}</p>{% endif %}
            {% if project.technologies %}<p><em>Technologies: {{ project.technologies | join(', ') }}</em></p>{% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    <div class="footer">
        <p><em>Generated on {{ export_date }}</em></p>
    </div>
</body>
</html>'''
            
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write(default_template)
